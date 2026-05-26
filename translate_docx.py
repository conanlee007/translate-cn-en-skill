#!/usr/bin/env python3
"""
DOCX Chinese → English translator with full format preservation.
Covers: paragraphs, table cells, headers/footers.

Usage: python translate_docx.py <input.docx> [output.docx]
Requires: ANTHROPIC_API_KEY environment variable
"""
import json
import os
import re
import sys
from pathlib import Path

import anthropic
from docx import Document
from docx.oxml.ns import qn

CLAUDE_MODEL = "claude-sonnet-4-6"
BATCH_SIZE = 20    # smaller batches → fewer JSON errors, faster per-call
MAX_WORKERS = 5    # concurrent API calls

CJK_RE = re.compile(r'[一-鿿㐀-䶿＀-￯　-〿⺀-⻿]')
FONT_NAME = 'Garamond'


def has_cjk(text: str) -> bool:
    return bool(CJK_RE.search(text))


# ── Claude translation (batch) ────────────────────────────────────────────────

import time
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed


def _parse_response(raw: str, expected_len: int, fallback: list[str]) -> list[str]:
    """
    Multi-strategy JSON parser — never raises, always returns a list of expected_len.
    Strategy 1: direct json.loads on extracted [...] block
    Strategy 2: regex extract all quoted strings
    Strategy 3: return fallback (original texts)
    """
    raw = re.sub(r'^```(?:json)?\s*', '', raw.strip())
    raw = re.sub(r'\s*```\s*$', '', raw)

    # Strategy 1: extract and parse JSON array
    m = re.search(r'\[.*\]', raw, re.DOTALL)
    if m:
        try:
            parsed = json.loads(m.group())
            if isinstance(parsed, list):
                result = [str(x) for x in parsed]
                if len(result) < expected_len:
                    result += fallback[len(result):]
                return result[:expected_len]
        except Exception:
            pass

    # Strategy 2: extract individual quoted strings via regex
    strings = re.findall(r'"((?:[^"\\]|\\.)*)"', raw)
    if strings:
        # Unescape
        strings = [s.replace('\\"', '"').replace('\\n', '\n').replace('\\\\', '\\')
                   for s in strings]
        if len(strings) < expected_len:
            strings += fallback[len(strings):]
        return strings[:expected_len]

    # Strategy 3: nothing worked, keep originals
    return fallback


def _api_call_with_backoff(client, texts: list[str]) -> str:
    """Call Claude API with exponential backoff on rate limit / server errors."""
    prompt = (
        "Translate the following Chinese texts to English. "
        "Return ONLY a JSON array with exactly the same number of elements in the same order.\n"
        "Rules:\n"
        "- Use standard English financial/accounting terminology\n"
        "- Express Chinese yuan as CNY (not RMB)\n"
        "- Keep numbers, percentages, dates, stock codes, and '--' exactly unchanged\n"
        "- Keep empty strings as empty strings\n"
        "- No explanations, no markdown fences — only the JSON array\n\n"
        + json.dumps(texts, ensure_ascii=False)
    )
    for attempt in range(5):
        try:
            msg = client.messages.create(
                model=CLAUDE_MODEL, max_tokens=8096,
                messages=[{"role": "user", "content": prompt}],
            )
            return msg.content[0].text
        except Exception as e:
            err = str(e).lower()
            if 'rate' in err or '429' in err or 'overload' in err:
                wait = 2 ** attempt * 10  # 10s, 20s, 40s, 80s, 160s
                print(f' [rate-limit, wait {wait}s]', end='', flush=True)
                time.sleep(wait)
            elif attempt < 4:
                time.sleep(3)
            else:
                raise
    raise RuntimeError("API call failed after 5 attempts")


def translate_batch(client, texts: list[str]) -> list[str]:
    """
    Translate a batch with progressive fallback:
      1. Full batch  →  2. Two halves  →  3. One-by-one
    JSON parsing uses multi-strategy parser, never crashes.
    """
    # Attempt 1: full batch
    try:
        raw = _api_call_with_backoff(client, texts)
        return _parse_response(raw, len(texts), texts)
    except Exception:
        pass

    # Attempt 2: two halves
    print(' [½-batch]', end='', flush=True)
    results = []
    mid = max(len(texts) // 2, 1)
    for chunk in [texts[:mid], texts[mid:]]:
        if not chunk:
            continue
        try:
            raw = _api_call_with_backoff(client, chunk)
            results.extend(_parse_response(raw, len(chunk), chunk))
        except Exception:
            # Attempt 3: one-by-one
            print(' [1-by-1]', end='', flush=True)
            for t in chunk:
                try:
                    raw = _api_call_with_backoff(client, [t])
                    results.extend(_parse_response(raw, 1, [t]))
                except Exception:
                    results.append(t)  # last resort: keep original
    return results


def translate_all(client, texts: list[str], checkpoint_path: str) -> list[str]:
    """
    Parallel batch translation with per-batch checkpointing.
    - BATCH_SIZE texts per API call (smaller = fewer JSON errors)
    - MAX_WORKERS concurrent threads
    - Checkpoint keyed by batch index → resume skips completed batches exactly
    - Thread-safe checkpoint writes via lock
    """
    total = len(texts)
    batches = [(idx, texts[i:i + BATCH_SIZE])
               for idx, i in enumerate(range(0, total, BATCH_SIZE))]
    n_batches = len(batches)

    # Load checkpoint (supports both old flat-list format and new dict format)
    completed: dict[int, list[str]] = {}
    if os.path.exists(checkpoint_path):
        try:
            with open(checkpoint_path, 'r', encoding='utf-8') as f:
                ckpt = json.load(f)
            if 'batches' in ckpt:
                completed = {int(k): v for k, v in ckpt['batches'].items()}
            elif 'results' in ckpt:
                # Migrate old format: reconstruct by batch index
                old = ckpt['results']
                for idx, i in enumerate(range(0, len(old), BATCH_SIZE)):
                    completed[idx] = old[i:i + BATCH_SIZE]
            done_texts = sum(len(v) for v in completed.values())
            print(f"    ↩ Resuming: {len(completed)}/{n_batches} batches done "
                  f"({done_texts}/{total} texts)")
        except Exception:
            completed = {}

    lock = threading.Lock()
    done_count = sum(len(v) for v in completed.values())

    def progress_line(done: int, batches_done: int) -> str:
        pct = done / total * 100
        bar_len = 30
        filled = int(bar_len * done / total)
        bar = '█' * filled + '░' * (bar_len - filled)
        return (f"  [{bar}] {done:,}/{total:,} texts "
                f"({pct:.1f}%) | batch {batches_done}/{n_batches}")

    def process_batch(idx: int, batch: list[str]) -> None:
        nonlocal done_count
        result = translate_batch(client, batch)
        with lock:
            completed[idx] = result
            done_count += len(result)
            batches_done = len(completed)
            print(progress_line(done_count, batches_done), flush=True)
            try:
                with open(checkpoint_path, 'w', encoding='utf-8') as f:
                    json.dump(
                        {'batches': {str(k): v for k, v in completed.items()}},
                        f, ensure_ascii=False
                    )
            except Exception as e:
                print(f'  [ckpt err: {e}]', flush=True)

    pending = [(idx, batch) for idx, batch in batches if idx not in completed]
    print(f"    {len(pending)} batches pending, {MAX_WORKERS} parallel workers")

    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        futures = {executor.submit(process_batch, idx, batch): idx
                   for idx, batch in pending}
        for future in as_completed(futures):
            try:
                future.result()
            except Exception as e:
                print(f"\n    [batch {futures[future]} error: {e}]")

    # Flatten in original order; fall back to original text for any missing batch
    results = []
    for idx, batch in batches:
        results.extend(completed.get(idx, batch))

    try:
        os.remove(checkpoint_path)
    except Exception:
        pass

    return results


# ── DOCX text extraction & writing ────────────────────────────────────────────

def para_full_text(para) -> str:
    return ''.join(run.text for run in para.runs)


def iter_header_footer_paras(doc: Document):
    for section in doc.sections:
        for hf in [
            section.header, section.footer,
            section.first_page_header, section.first_page_footer,
            section.even_page_header, section.even_page_footer,
        ]:
            if hf is not None:
                yield from hf.paragraphs
                for table in hf.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            yield from cell.paragraphs


def collect_paragraph_texts(doc: Document):
    refs, texts = [], []
    for para in list(doc.paragraphs) + list(iter_header_footer_paras(doc)):
        full = para_full_text(para)
        if has_cjk(full):
            refs.append(para)
            texts.append(full)
    return refs, texts


def collect_table_texts(doc: Document):
    refs, texts = [], []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full = para_full_text(para)
                    if has_cjk(full):
                        refs.append(para)
                        texts.append(full)
    return refs, texts


def set_run_font(run, font_name: str):
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'))
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), font_name)


def set_para_text(para, new_text: str):
    if not para.runs:
        return
    para.runs[0].text = new_text
    set_run_font(para.runs[0], FONT_NAME)
    for run in para.runs[1:]:
        run.text = ''


# ── Main ──────────────────────────────────────────────────────────────────────

def translate_docx(input_path: str, output_path: str):
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        sys.exit('Error: ANTHROPIC_API_KEY is not set.')

    client = anthropic.Anthropic(api_key=api_key)
    doc = Document(input_path)

    # 1. Text blocks
    print('Collecting text from paragraphs & headers/footers...')
    para_refs, para_texts = collect_paragraph_texts(doc)
    print(f'  {len(para_texts)} paragraphs with Chinese text')

    print('Collecting text from tables...')
    cell_refs, cell_texts = collect_table_texts(doc)
    print(f'  {len(cell_texts)} table cell paragraphs with Chinese text')

    all_refs = para_refs + cell_refs
    all_texts = para_texts + cell_texts
    checkpoint = output_path + '.checkpoint.json'
    print(f'\nTranslating {len(all_texts)} text blocks...')
    all_translated = translate_all(client, all_texts, checkpoint)

    print('\nWriting text translations back...')
    for para, translated in zip(all_refs, all_translated):
        set_para_text(para, translated)

    doc.save(output_path)
    print(f'\nSaved → {output_path}')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Usage: python translate_docx.py <input.docx> [output.docx]')
        sys.exit(1)
    inp = sys.argv[1]
    out = (sys.argv[2] if len(sys.argv) > 2
           else str(Path(inp).parent / (Path(inp).stem + '-EN.docx')))
    translate_docx(inp, out)
